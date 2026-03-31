import os
import pandas as pd
import PyPDF2
from docx import Document
import json
import xml.etree.ElementTree as ET
from typing import Dict, Any, Optional
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor:
    """Universal file content extractor for all file types"""
    
    def __init__(self):
        self.supported_extensions = {
            'txt', 'md', 'json', 'xml', 'csv', 'xlsx', 'xls',
            'pdf', 'docx', 'doc'
        }
    
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from any supported file type
        Returns a dictionary with content, metadata, and processing info
        """
        if not os.path.exists(file_path):
            return {"error": "File not found", "content": "", "metadata": {}}
        
        file_extension = self._get_file_extension(file_path)
        
        if file_extension not in self.supported_extensions:
            return {
                "error": f"Unsupported file type: {file_extension}",
                "content": "",
                "metadata": {"file_type": file_extension}
            }
        
        try:
            if file_extension in ['txt', 'md']:
                return self._extract_text_file(file_path)
            elif file_extension == 'json':
                return self._extract_json_file(file_path)
            elif file_extension == 'xml':
                return self._extract_xml_file(file_path)
            elif file_extension in ['csv', 'xlsx', 'xls']:
                return self._extract_spreadsheet(file_path)
            elif file_extension == 'pdf':
                return self._extract_pdf_file(file_path)
            elif file_extension in ['docx', 'doc']:
                return self._extract_word_document(file_path)
            else:
                return {"error": f"Unsupported file type: {file_extension}", "content": "", "metadata": {}}
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return {
                "error": f"Error processing file: {str(e)}",
                "content": "",
                "metadata": {"file_type": file_extension}
            }
    
    def _get_file_extension(self, file_path: str) -> str:
        """Get file extension in lowercase"""
        return file_path.split('.')[-1].lower()
    
    def _extract_text_file(self, file_path: str) -> Dict[str, Any]:
        """Extract content from text files"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                # Try with different encoding
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
            except Exception as e:
                content = f"Error reading file: {str(e)}"
        
        return {
            "content": content,
            "metadata": {
                "file_type": "text",
                "encoding": "utf-8",
                "lines": len(content.split('\n')),
                "characters": len(content)
            }
        }
    
    def _extract_json_file(self, file_path: str) -> Dict[str, Any]:
        """Extract content from JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                content = json.dumps(data, indent=2)
        except Exception as e:
            content = f"Error parsing JSON: {str(e)}"
        
        return {
            "content": content,
            "metadata": {
                "file_type": "json",
                "parsed": True
            }
        }
    
    def _extract_xml_file(self, file_path: str) -> Dict[str, Any]:
        """Extract content from XML files"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            content = ET.tostring(root, encoding='unicode', method='xml')
        except Exception as e:
            content = f"Error parsing XML: {str(e)}"
        
        return {
            "content": content,
            "metadata": {
                "file_type": "xml",
                "parsed": True
            }
        }
    
    def _extract_spreadsheet(self, file_path: str) -> Dict[str, Any]:
        """Extract content from spreadsheet files"""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Generate basic statistics
            num_rows = len(df)
            num_cols = len(df.columns)
            numeric_cols = df.select_dtypes(include=['int64', 'float64']).columns
            
            # Create a concise data preview
            preview_rows = min(5, num_rows)
            preview = df.head(preview_rows).to_string()
            
            # Add summary information
            summary = f"Spreadsheet Analysis:\n"
            summary += f"- Total Rows: {num_rows}\n"
            summary += f"- Total Columns: {num_cols}\n"
            summary += f"- Column Names: {', '.join(df.columns)}\n"
            
            if numeric_cols.any():
                summary += "\nNumeric Column Statistics:\n"
                for col in numeric_cols:
                    stats = df[col].describe()
                    summary += f"  {col}:\n"
                    summary += f"    - Mean: {stats['mean']:.2f}\n"
                    summary += f"    - Min: {stats['min']:.2f}\n"
                    summary += f"    - Max: {stats['max']:.2f}\n"
            
            summary += f"\nFirst {preview_rows} rows of data:\n{preview}"
            
            # For large datasets, indicate truncation
            if num_rows > preview_rows:
                summary += f"\n\n(Showing {preview_rows} of {num_rows} total rows)"
            
            return {
                "content": summary,
                "metadata": {
                    "file_type": "spreadsheet",
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": list(df.columns)
                }
            }
        except Exception as e:
            return {
                "content": f"Error reading spreadsheet: {str(e)}",
                "metadata": {"file_type": "spreadsheet"}
            }
    
    def _extract_pdf_file(self, file_path: str) -> Dict[str, Any]:
        """Extract content from PDF files"""
        try:
            import re
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                content = ""
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Post-process: join words that were split across lines
                        # PyPDF2 often returns each word on a new line - need to rejoin them
                        lines = page_text.split('\n')
                        cleaned_parts = []
                        for line in lines:
                            line = line.strip()
                            if line:
                                cleaned_parts.append(line)
                        
                        # Join all parts with spaces first
                        joined = ' '.join(cleaned_parts)
                        
                        # Split on sentence endings to reconstruct paragraphs
                        # Split on: period + space + capital, or question/exclamation + space + capital
                        parts = re.split(r'([.!?]\s+[A-Z])', joined)
                        
                        # Reconstruct text: join each delimiter with the following text
                        page_text = ""
                        for i, part in enumerate(parts):
                            if i % 2 == 0:
                                # Regular text
                                page_text += part
                            else:
                                # Captured delimiter (punctuation+space+capital) - keep it
                                page_text += part
                        
                        # If reconstruction failed or is empty, use joined version
                        if not page_text.strip():
                            page_text = joined
                        
                        content += f"\n--- Page {page_num} ---\n{page_text}\n"
                
                if not content.strip():
                    content = "No text content found in PDF (may be image-based)"
                
                return {
                    "content": content,
                    "metadata": {
                        "file_type": "pdf",
                        "pages": page_count,
                        "has_text": bool(content.strip())
                    }
                }
        except Exception as e:
            return {
                "content": f"Error reading PDF: {str(e)}",
                "metadata": {"file_type": "pdf"}
            }
    
    def _extract_word_document(self, file_path: str) -> Dict[str, Any]:
        """Extract content from Word documents"""
        try:
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            return {
                "content": content,
                "metadata": {
                    "file_type": "word_document",
                    "paragraphs": len(doc.paragraphs)
                }
            }
        except Exception as e:
            return {
                "content": f"Error reading Word document: {str(e)}",
                "metadata": {"file_type": "word_document"}
            }

# Global instance
file_processor = FileProcessor() 